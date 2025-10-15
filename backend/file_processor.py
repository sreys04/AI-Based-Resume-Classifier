"""
File processing utilities for extracting text from resume documents
"""
import os
import PyPDF2
import pdfplumber
from docx import Document
import re

def clean_text(text):
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove extra whitespace and normalize
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep important ones
    text = re.sub(r'[^\w\s@.,()-]', ' ', text)
    
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def extract_text_from_pdf(file_path):
    """Extract text from PDF file using multiple methods"""
    text = ""
    
    # Method 1: Try pdfplumber first (better for complex layouts)
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            return clean_text(text)
    except Exception as e:
        print(f"pdfplumber failed: {e}")
    
    # Method 2: Fallback to PyPDF2
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        return clean_text(text)
    except Exception as e:
        print(f"PyPDF2 failed: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return clean_text(text)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return clean_text(text)
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            return clean_text(text)
        except Exception as e:
            print(f"TXT extraction failed: {e}")
            return ""
    except Exception as e:
        print(f"TXT extraction failed: {e}")
        return ""

def process_resume_file(file_path):
    """
    Main function to process resume file and extract text
    
    Args:
        file_path (str): Path to the resume file
        
    Returns:
        str: Extracted and cleaned text content
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Get file extension
    file_extension = os.path.splitext(file_path)[1].lower()
    
    # Process based on file type
    if file_extension == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        text = extract_text_from_docx(file_path)
    elif file_extension in ['.txt', '.text']:
        text = extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")
    
    # Validate extracted text
    if not text or len(text.strip()) < 50:
        raise ValueError("Insufficient text extracted from document. Please ensure the file contains readable text.")
    
    return text

def get_file_info(file_path):
    """Get basic information about a file"""
    if not os.path.exists(file_path):
        return None
    
    stat = os.stat(file_path)
    file_extension = os.path.splitext(file_path)[1].lower()
    
    return {
        'size': stat.st_size,
        'extension': file_extension,
        'type': get_file_type(file_extension),
        'readable': file_extension in ['.pdf', '.docx', '.txt', '.text']
    }

def get_file_type(extension):
    """Map file extension to human-readable type"""
    type_map = {
        '.pdf': 'PDF Document',
        '.docx': 'Word Document',
        '.txt': 'Text File',
        '.text': 'Text File'
    }
    return type_map.get(extension, 'Unknown')